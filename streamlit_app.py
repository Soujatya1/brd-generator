import streamlit as st
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from docx import Document
from io import BytesIO
import hashlib
import PyPDF2
import os
import pdfplumber

# Initialize the model
model = ChatGroq(
    groq_api_key="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", 
    model_name="Llama3-70b-8192"
)

llm_chain = LLMChain(llm=model, prompt=PromptTemplate(
    input_variables=['template_format', 'requirements', 'tables'],
    template="""Create a Business Requirements Document (BRD) based on the following details:

    Document Structure:
    {template_format}

    Requirements:
    Extract and organize the content from the provided documents and map each piece of information to the corresponding sections in the BRD. Be specific and avoid extraneous information.

    Tables:
    If applicable, include the following tabular information extracted from the documents. Ensure correct formatting and alignment:
    {tables}

    Formatting:
    - Use clear headings, subheadings, and numbered bullet points for organization.
    - Differentiate between functional and non-functional requirements clearly.
    - Provide tables in a neat and structured format.
    Also, if necssary insert a fields from the 
    Output:
    The result must strictly follow the BRD format specified above, and be concise yet comprehensive. Avoid redundancy and unnecessary speculation.
    """
))

# Function to extract text and tables from .docx files
def extract_content_from_docx(file):
    doc = Document(file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            table_data.append([cell.text.strip() for cell in row.cells])
        tables.append(table_data)
    # Convert tables to a clean text format (tabular format for LLM processing)
    tables_as_text = "\n\n".join(["\n".join(["\t".join(row) for row in table]) for table in tables])
    return text, tables_as_text

# Function to extract text and tables from .pdf files
def extract_text_from_pdf(file):
    text = ""
    tables = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text()
            # Try extracting tables as well
            for table in page.extract_tables():
                tables += "\n" + "\n".join(["\t".join(row) for row in table])
    return text, tables

# Streamlit UI setup
st.title("BRD Generator")
st.write("Upload requirement documents and define the BRD structure below to generate a detailed Business Requirements Document.")

uploaded_files = st.file_uploader("Upload requirement documents (PDF/DOCX):", accept_multiple_files=True)
template_format = st.text_area("Enter the BRD format:", height=200, placeholder="Define the structure of the BRD here...")

# Processing uploaded files
if uploaded_files:
    combined_requirements = ""
    all_tables_as_text = ""
    for uploaded_file in uploaded_files:
        file_extension = os.path.splitext(uploaded_file.name)[-1].lower()
        if file_extension == ".docx":
            text, tables_as_text = extract_content_from_docx(uploaded_file)
            combined_requirements += text + "\n"
            all_tables_as_text += tables_as_text + "\n"
        elif file_extension == ".pdf":
            text, tables_as_text = extract_text_from_pdf(uploaded_file)
            combined_requirements += text + "\n"
            all_tables_as_text += tables_as_text + "\n"
        else:
            st.warning(f"Unsupported file format: {uploaded_file.name}")
    
    requirements = combined_requirements
else:
    requirements = ""
    all_tables_as_text = ""

# Hash function for caching
def generate_hash(template_format, requirements):
    combined_string = template_format + requirements
    return hashlib.md5(combined_string.encode()).hexdigest()

if "outputs_cache" not in st.session_state:
    st.session_state.outputs_cache = {}

# Generate BRD when button is clicked
if st.button("Generate BRD") and requirements and template_format:
    doc_hash = generate_hash(template_format, requirements)
    
    if doc_hash in st.session_state.outputs_cache:
        output = st.session_state.outputs_cache[doc_hash]
    else:
        prompt_input = {
            "template_format": template_format,
            "requirements": requirements,
            "tables": all_tables_as_text,
        }
        output = llm_chain.run(prompt_input)
        st.session_state.outputs_cache[doc_hash] = output
    st.write(output)

    # Create a Word document from the BRD
    doc = Document()
    doc.add_heading('Business Requirements Document', level=1)
    doc.add_paragraph(output, style='Normal')

    # Add the tables to the Word document
    if all_tables_as_text:
        doc.add_heading("Tables", level=2)
        doc.add_paragraph(all_tables_as_text, style='Normal')

    # Save the Word document to a buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Provide a download link for the BRD document
    st.download_button(
        label="Download BRD as Word document",
        data=buffer,
        file_name="BRD.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

output = st.session_state.outputs_cache

if isinstance(output, dict) and 'text' in output:
    output_text = output['text']
else:
    output_text = str(output)
    
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import difflib

# Function to calculate text similarity using Cosine Similarity
def calculate_text_similarity(text1, text2):
    vectorizer = TfidfVectorizer().fit_transform([text1, text2])
    vectors = vectorizer.toarray()
    cosine_sim = cosine_similarity(vectors)
    return cosine_sim[0][1] * 100  # Return as a percentage

# Function to calculate structural similarity
def calculate_structural_similarity(tables1, tables2):
    sm = difflib.SequenceMatcher(None, tables1, tables2)
    return sm.ratio() * 100  # Return as a percentage

# New section: Upload a sample document for comparison
st.write("Optional: Upload a sample BRD for comparison.")
sample_file = st.file_uploader("Upload a sample BRD (PDF/DOCX):", type=["pdf", "docx"])

if sample_file:
    file_extension = os.path.splitext(sample_file.name)[-1].lower()
    if file_extension == ".docx":
        sample_text, sample_tables = extract_content_from_docx(sample_file)
    elif file_extension == ".pdf":
        sample_text = extract_text_from_pdf(sample_file)
        sample_tables = ""  # No table support in the current PDF extraction function
    else:
        st.warning(f"Unsupported file format: {sample_file.name}")
        sample_text, sample_tables = "", ""

    # Ensure there's content to compare
    if requirements and template_format and sample_text:
        # Calculate match scores
        content_similarity = calculate_text_similarity(requirements, sample_text)
        content_similarity_1 = calculate_text_similarity(output_text, sample_text)
        format_similarity = calculate_structural_similarity(all_tables_as_text, sample_tables)
        st.subheader("Match Score Results")
        if content_similarity_1 != 0:
            similarity_ratio = content_similarity_1 / content_similarity
            st.write(f"Similarity Ratio: {similarity_ratio:.2f}")
        else:
            print("Error: content_similarity_1 is 0, division by zero is not possible.")
        # Final weighted score
        content_weight = 0.7
        format_weight = 0.3
        final_score = (content_similarity * content_weight) + (format_similarity * format_weight)

        # Display results
        #st.subheader("Match Score Results")
        #st.write(f"Content Match: {content_similarity:.2f}%")
        #st.write(f"Format Match: {format_similarity:.2f}%")
        #st.write(f"Overall Match Score: {final_score:.2f}%")
    else:
        st.warning("Please generate a BRD first or ensure the sample document has valid content.")
